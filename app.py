import os
import io
import json
import re
import google.generativeai as genai
import docx
from flask import Flask, request, render_template_string, send_file, flash, redirect, url_for
from dotenv import load_dotenv
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT

# --- CONFIGURATION ---
load_dotenv() # Load environment variables from .env file

app = Flask(__name__)
app.secret_key = 'super_secret_key_for_flask_flashing' # Needed for flash messages

# Configure Gemini API
try:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY not found in .env file.")
    genai.configure(api_key=api_key)
    GEMINI_CONFIGURED = True
except Exception as e:
    print(f"⚠️ API Key Configuration Error: {e}")
    GEMINI_CONFIGURED = False

# --- HTML TEMPLATE ---
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Student Document Generator</title>
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 40px auto; padding: 20px; background-color: #f4f7f9; }
        .container { background-color: #fff; border-radius: 8px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); padding: 30px; }
        h1 { color: #2c3e50; text-align: center; }
        label { font-weight: bold; color: #34495e; display: block; margin-bottom: 8px; }
        textarea, input[type="file"] { width: 100%; padding: 10px; border-radius: 4px; border: 1px solid #ccc; margin-bottom: 20px; box-sizing: border-box; }
        button { background-color: #3498db; color: white; padding: 12px 20px; border: none; border-radius: 4px; cursor: pointer; font-size: 16px; width: 100%; transition: background-color 0.3s; }
        button:hover { background-color: #2980b9; }
        .flash { padding: 15px; margin-bottom: 20px; border-radius: 4px; text-align: center; }
        .flash.error { background-color: #e74c3c; color: white; }
        .flash.success { background-color: #2ecc71; color: white; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 Student Document Generator</h1>
        
        {% with messages = get_flashed_messages(with_categories=true) %}
          {% if messages %}
            {% for category, message in messages %}
              <div class="flash {{ category }}">{{ message }}</div>
            {% endfor %}
          {% endif %}
        {% endwith %}

        <form method="post" enctype="multipart/form-data">
            <label for="file">Upload your raw text file (.txt):</label>
            <input type="file" id="file" name="file" accept=".txt" required>
            <button type="submit">Generate Document</button>
        </form>
    </div>
</body>
</html>
"""

# --- LOGIC FROM json_generator.py ---

def create_gemini_prompt(unstructured_data):
    """Creates the final, ultra-specific prompt with mandatory fields."""
    return f"""
You are a highly precise data structuring engine. Your task is to convert the provided text into a single, clean JSON array. You must follow the rules below without exception.
**CRITICAL RULES:**
1.  **Mandatory Fields:** Every single student object in the output array **MUST** contain these four keys: `full_name`, `registration_number`, `department`, and `year`. No exceptions.
2.  **Handling Missing Information:** If the information for any of the four mandatory keys cannot be found or inferred from the text, you **MUST** include the key with an **empty string `""`** as its value. Do not omit the key.
3.  **Department Inference (Mandatory):**
    * First, try to extract the department from the text (e.g., "ECE", "CSE").
    * If it's not in the text, you **MUST** infer it from the two-letter code in the `registration_number`.
    * **Department Code Mapping:**
      `CS`→`CSE`, `IT`→`IT`, `EC`→`ECE`, `EE`→`EEE`, `CE`→`CIVIL`, `ME`→`MECH`, `AD`→`AI&DS`, `AM`→`AIML`, `EI`→`EIE`, `CB`→`CSBS`, `CJ`→`M.Tech CSE`, `MU`→`Mechanical and Automation`, `IC`→`ICE`.
    * If the department cannot be found or inferred, set its value to an empty string: `"department": ""`.
4.  **Year Inference (Mandatory):**
    * You **MUST** infer the `year` from the `registration_number`.
    * Reference the current academic year: 2025-2026.
    * **Inference Logic:**
        `SEC22...` → `"year": "Fourth"`
        `SEC23...` → `"year": "Third"`
        `SEC24...` → `"year": "Second"`
        `SEC25...` → `"year": "First"`
    * If the `registration_number` is missing, set the year's value to an empty string: `"year": ""`.
5.  **Category Key (Conditional):**
    * The `category` key is the **ONLY** key you should omit.
    * Only include `category: "Hostel"` if the student is explicitly listed under a "(Hostellers Only)" section.
    * Otherwise, the `category` key **MUST NOT** be present in the object.
6.  **Final Output Format:** Your entire response **MUST** be only the JSON array `[...]`. Do not include any text or markdown before or after it.
---
**Now, process the following input data according to these exact rules:**
{unstructured_data}
"""

def extract_json_from_response(text):
    """Extracts a JSON array from the model's response, handling markdown code fences."""
    match = re.search(r'\[.*\]', text, re.DOTALL)
    if match:
        return match.group(0)
    return None

# --- EXACT LOGIC FROM od_program.py ---

def add_senior_student_list(doc, students):
    """Adds a formatted, numbered list of senior students to the document with a centered hyphen."""
    for i, student in enumerate(students, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.25)
        pf.tab_stops.clear_all()
        # Left-aligned tab for student name
        pf.tab_stops.add_tab_stop(Inches(0.5))
        # Centered tab for hyphen
        pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.CENTER)
        # Right-aligned tab for registration number
        pf.tab_stops.add_tab_stop(Inches(4.0))

        # Add runs separately
        run1 = p.add_run(f"{i}.\t{student['full_name']}\t")
        run1.font.color.rgb = RGBColor(0, 0, 0)

        run2 = p.add_run("–")
        run2.font.color.rgb = RGBColor(0, 0, 0)

        run3 = p.add_run(f"\t{student['registration_number']}")
        run3.font.color.rgb = RGBColor(0, 0, 0)

def add_first_year_table(doc, students):
    """Adds a formatted table of first-year students to the document."""
    if not students:
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Disable autofit and set explicit column widths to control layout
    table.autofit = False
    table.allow_autofit = False

    table.columns[0].width = Inches(0.1)  # S.No. (Narrow)
    table.columns[1].width = Inches(2.0)  # Name (Wider)
    table.columns[2].width = Inches(1.0)  # SEC ID
    table.columns[3].width = Inches(0.1)  # Section (Narrow)
    table.columns[4].width = Inches(1.0)  # Department

    hdr_cells = table.rows[0].cells
    headers = ['S.No.', 'Name', 'SEC ID', 'Section', 'Department']

    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Group students by department before adding to the table
    sorted_students = sorted(students, key=lambda s: s['department'])
    for i, student in enumerate(sorted_students, 1):
        row_cells = table.add_row().cells
        row_data = [
            str(i),
            student['full_name'],
            student['registration_number'],
            student.get('section', 'N/A'),
            student['department']
        ]
        for idx, cell_text in enumerate(row_data):
            p = row_cells[idx].paragraphs[0]
            p.add_run(cell_text).font.color.rgb = RGBColor(0, 0, 0)

def generate_student_document(students):
    """
    Generates a Word document with a Category -> Year -> Department hierarchy.
    """
    doc = docx.Document()
    style = doc.styles['Normal'].font
    style.name = 'Times New Roman'
    style.size = Pt(12)

    # Changed the title from "Student Details" to "VOLUNTEERS LIST"
    doc.add_heading("VOLUNTEERS LIST", 0).runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Step 1: Group all students by category
    categorized_students = {}
    for student in students:
        # Default category is "Uncategorized" if not present in the student data
        category = student.get("category", "Uncategorized")
        if category not in categorized_students:
            categorized_students[category] = []
        categorized_students[category].append(student)

    # Define the order for categories
    category_order = ["Hostel", "Dayscholar", "Uncategorized"]

    for category_name in category_order:
        student_list = categorized_students.get(category_name)
        if not student_list:
            continue

        # Add Category Heading (Level 1)
        if category_name=="Uncategorized":
            pass
        else:
            doc.add_heading(category_name, level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Step 2: Group students within the category by year
        students_by_year = {}
        for student in student_list:
            year = student["year"]
            if year not in students_by_year:
                students_by_year[year] = []
            students_by_year[year].append(student)

        year_order = ["Fourth", "Third", "Second", "First"]
        for year in year_order:
            year_student_list = students_by_year.get(year)
            if not year_student_list:
                continue

            # Add Year Heading (Level 2)
            doc.add_heading(f"{year} Year", level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)

            # Step 3: Handle First Years and Seniors differently
            if year == "First":
                # For First Years, add the table directly under the year heading
                add_first_year_table(doc, year_student_list)
            else:
                # For Seniors, group them by department
                students_by_dept = {}
                for student in year_student_list:
                    dept = student["department"]
                    if dept not in students_by_dept:
                        students_by_dept[dept] = []
                    students_by_dept[dept].append(student)

                for dept_name in sorted(students_by_dept.keys()):
                    dept_student_list = students_by_dept[dept_name]
                    # Add Department Heading (Level 3)
                    doc.add_heading(dept_name, level=3).runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # Add the numbered list of students
                    add_senior_student_list(doc, dept_student_list)

    # NECESSARY MODIFICATION FOR WEB APP: Save to an in-memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning
    return file_stream

# --- FLASK WEB ROUTES ---

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if not GEMINI_CONFIGURED:
            flash("Error: Gemini API key is not configured. Please check your .env file and restart the server.", "error")
            return redirect(url_for('index'))
            
        if 'file' not in request.files:
            flash("Error: No file part in the request.", "error")
            return redirect(url_for('index'))
        file = request.files['file']
        if file.filename == '':
            flash("Error: No file selected.", "error")
            return redirect(url_for('index'))
            
        if file and file.filename.endswith('.txt'):
            try:
                # --- Step A: Process with Gemini ---
                unstructured_data = file.read().decode("utf-8")
                if not unstructured_data.strip():
                    flash("Error: The uploaded file is empty.", "error")
                    return redirect(url_for('index'))

                model = genai.GenerativeModel("models/gemini-pro-latest")
                prompt = create_gemini_prompt(unstructured_data)
                response = model.generate_content(prompt)
                
                json_string = extract_json_from_response(response.text)
                if not json_string:
                    raise ValueError("Gemini did not return a valid JSON array.")
                
                student_data = json.loads(json_string)

                # --- Step B: Generate the DOCX ---
                document_stream = generate_student_document(student_data)
                
                # --- Step C: Send the file for download ---
                return send_file(
                    document_stream,
                    as_attachment=True,
                    download_name='student_details.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            except Exception as e:
                flash(f"An error occurred: {e}", "error")
                return redirect(url_for('index'))

    # For GET request, just show the upload page
    return render_template_string(HTML_TEMPLATE)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)